from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_with_style(doc, text, level):
    """Add heading with custom styling"""
    heading = doc.add_heading(text, level)
    return heading

def create_table_with_header(doc, headers, rows, col_widths=None):
    """Create a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], '2E86AB')
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for row_data in rows:
        row = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row[i].text = str(cell_data)

    return table

# Create document
doc = Document()

# Title
title = doc.add_heading('SoundPlus++ Project', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('CI/CD Pipeline Design and Automation Documentation')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(14)
    run.font.italic = True

# Project info
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Premium Audio Equipment E-commerce Platform\n').bold = True
info.add_run('MERN Stack Application with Docker Containerization\n')
info.add_run('Version 1.0 | January 2026')

doc.add_page_break()

# Table of Contents
doc.add_heading('Table of Contents', 1)

toc_items = [
    ('1. Introduction', '3'),
    ('   1.1 Project Overview', '3'),
    ('   1.2 Technology Stack Summary', '3'),
    ('2. Part 1: CI/CD Design Diagram', '4'),
    ('   2.1 Architecture Overview Diagram', '4'),
    ('   2.2 CI/CD Pipeline Flow Diagram', '5'),
    ('   2.3 Container Architecture Diagram', '6'),
    ('   2.4 Component Connectivity Diagram', '7'),
    ('   2.5 Diagram Explanation', '8'),
    ('3. Part 2: Automation Approach', '10'),
    ('   3.1 DevOps Tools and Versions', '10'),
    ('   3.2 Application Tools and Dependencies', '11'),
    ('   3.3 Jenkins Pipeline Stages', '13'),
    ('   3.4 GitHub Actions Pipeline', '14'),
    ('   3.5 Deployment Automation Flow', '15'),
    ('4. Environment Configuration', '16'),
    ('5. Security Considerations', '17'),
    ('6. Conclusion', '18'),
]

for item, page in toc_items:
    p = doc.add_paragraph()
    p.add_run(item)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6))
    p.add_run('\t' + page)

doc.add_page_break()

# Section 1: Introduction
doc.add_heading('1. Introduction', 1)

doc.add_heading('1.1 Project Overview', 2)
doc.add_paragraph(
    'SoundPlus++ is a premium audio equipment e-commerce platform built using the MERN stack '
    '(MongoDB, Express.js, React, Node.js). The application provides a comprehensive online '
    'shopping experience for audio enthusiasts, featuring product catalog management, user '
    'authentication, shopping cart functionality, and order processing.'
)

doc.add_paragraph(
    'This document outlines the CI/CD (Continuous Integration/Continuous Deployment) design '
    'and automation approach implemented for the SoundPlus++ application, ensuring reliable '
    'and efficient software delivery.'
)

doc.add_heading('1.2 Technology Stack Summary', 2)

stack_data = [
    ('Layer', 'Technology', 'Version'),
]
stack_rows = [
    ('Frontend Framework', 'React', '18.3.1'),
    ('Build Tool', 'Vite', '6.0.5'),
    ('Backend Framework', 'Express.js', '4.21.2'),
    ('Runtime', 'Node.js', '18.x (LTS)'),
    ('Database', 'MongoDB Atlas', 'Cloud'),
    ('Containerization', 'Docker', 'Latest'),
    ('Orchestration', 'Docker Compose', 'v2'),
    ('CI/CD (Local)', 'Jenkins', 'LTS'),
    ('CI/CD (Cloud)', 'GitHub Actions', 'Latest'),
    ('Version Control', 'Git/GitHub', 'Latest'),
]

create_table_with_header(doc, ['Layer', 'Technology', 'Version'], stack_rows)

doc.add_page_break()

# Section 2: Part 1 - CI/CD Design Diagram
doc.add_heading('2. Part 1: CI/CD Design Diagram', 1)

doc.add_heading('2.1 Architecture Overview Diagram', 2)

doc.add_paragraph(
    'The following diagram illustrates the complete CI/CD architecture for the SoundPlus++ '
    'application, showing all major components and their interconnections.'
)

# ASCII Diagram 1: Architecture Overview
diagram1 = doc.add_paragraph()
diagram1.alignment = WD_ALIGN_PARAGRAPH.CENTER
diagram1_text = '''
+============================================================================+
|                    SOUNDPLUS++ CI/CD ARCHITECTURE                          |
+============================================================================+

    +---------------+          +------------------+          +---------------+
    |   DEVELOPER   |  push    |     GITHUB       |  webhook |    JENKINS    |
    |   Workstation |--------->|   Repository     |--------->|    Server     |
    +---------------+          +------------------+          +---------------+
          |                           |                            |
          |                           | trigger                    | build
          v                           v                            v
    +---------------+          +------------------+          +---------------+
    |   Local Dev   |          |  GitHub Actions  |          |   Docker      |
    |   Environment |          |   CI Pipeline    |          |   Build       |
    +---------------+          +------------------+          +---------------+
                                      |                            |
                                      | push                       | push
                                      v                            v
                               +------------------+          +---------------+
                               |   DOCKER HUB     |<---------|   Docker      |
                               |   Registry       |          |   Images      |
                               +------------------+          +---------------+
                                      |
                                      | pull
                                      v
                   +------------------------------------------+
                   |           DEPLOYMENT TARGET              |
                   |  +----------------+  +----------------+  |
                   |  |   AWS EC2      |  |   Local VMs    |  |
                   |  |   Instance     |  |   (Docker)     |  |
                   |  +----------------+  +----------------+  |
                   +------------------------------------------+
                                      |
                                      v
                   +------------------------------------------+
                   |              DOCKER HOST                 |
                   |  +----------------+  +----------------+  |
                   |  |   Frontend     |  |   Backend      |  |
                   |  |   Container    |  |   Container    |  |
                   |  |   (Port 3000)  |  |   (Port 5000)  |  |
                   |  +----------------+  +----------------+  |
                   |             |              |             |
                   |             +------+-------+             |
                   |                    v                     |
                   |           +----------------+             |
                   |           | MongoDB Atlas  |             |
                   |           |   (Cloud DB)   |             |
                   |           +----------------+             |
                   +------------------------------------------+
'''
for run in diagram1.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
diagram1.add_run(diagram1_text)
for run in diagram1.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

doc.add_paragraph('Figure 2.1: SoundPlus++ CI/CD Architecture Overview')

doc.add_page_break()

doc.add_heading('2.2 CI/CD Pipeline Flow Diagram', 2)

doc.add_paragraph(
    'This diagram shows the detailed flow of the CI/CD pipeline from code commit to deployment.'
)

diagram2 = doc.add_paragraph()
diagram2.alignment = WD_ALIGN_PARAGRAPH.CENTER
diagram2_text = '''
+============================================================================+
|                      CI/CD PIPELINE FLOW DIAGRAM                           |
+============================================================================+

  [Developer]
       |
       | git push
       v
  +----------+     +----------+     +----------+     +----------+
  |  COMMIT  |---->|  GITHUB  |---->| WEBHOOK  |---->| JENKINS  |
  |   Code   |     |   Repo   |     | Trigger  |     |  Server  |
  +----------+     +----------+     +----------+     +----------+
                                                           |
       +---------------------------------------------------+
       |
       v
  +============================================================================+
  |                         JENKINS PIPELINE STAGES                            |
  +============================================================================+
  |                                                                            |
  |  Stage 1         Stage 2          Stage 3         Stage 4                  |
  |  +---------+     +-----------+    +----------+    +------------+           |
  |  |CHECKOUT |---->|PRE-FLIGHT |---->| SETUP   |---->|   BUILD   |           |
  |  |  Code   |     |  Check    |    |   ENV    |    |  Images   |           |
  |  +---------+     +-----------+    +----------+    +------------+           |
  |                                                          |                 |
  |                                                          v                 |
  |  Stage 7         Stage 6          Stage 5                                  |
  |  +---------+     +-----------+    +------------+                           |
  |  | SUCCESS |<----|  VERIFY   |<---|   START    |                           |
  |  | Report  |     | Services  |    |  Services  |                           |
  |  +---------+     +-----------+    +------------+                           |
  |                                                                            |
  +============================================================================+
       |
       v
  +----------+     +----------+     +----------+
  | DOCKER   |---->| HEALTH   |---->| DEPLOY   |
  | Registry |     |  Check   |     | Complete |
  +----------+     +----------+     +----------+
'''
diagram2.add_run(diagram2_text)
for run in diagram2.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

doc.add_paragraph('Figure 2.2: CI/CD Pipeline Flow')

doc.add_page_break()

doc.add_heading('2.3 Container Architecture Diagram', 2)

doc.add_paragraph(
    'The following diagram illustrates the Docker container architecture and internal '
    'connectivity of the SoundPlus++ application.'
)

diagram3 = doc.add_paragraph()
diagram3.alignment = WD_ALIGN_PARAGRAPH.CENTER
diagram3_text = '''
+============================================================================+
|                   CONTAINER ARCHITECTURE DIAGRAM                           |
+============================================================================+

                            DOCKER HOST
+------------------------------------------------------------------------+
|                                                                        |
|   soundplus-network (bridge)                                           |
|   +----------------------------------------------------------------+   |
|   |                                                                |   |
|   |  +------------------------+      +------------------------+    |   |
|   |  |  soundplus-frontend    |      |  soundplus-backend     |    |   |
|   |  |  Container             |      |  Container             |    |   |
|   |  +------------------------+      +------------------------+    |   |
|   |  |                        |      |                        |    |   |
|   |  |  +------------------+  |      |  +------------------+  |    |   |
|   |  |  |   React App      |  |      |  |   Express.js     |  |    |   |
|   |  |  |   (Vite Dev)     |  | HTTP |  |   REST API       |  |    |   |
|   |  |  |                  |<-|------|->|                  |  |    |   |
|   |  |  |   Port: 3000     |  |      |  |   Port: 5000     |  |    |   |
|   |  |  +------------------+  |      |  +------------------+  |    |   |
|   |  |                        |      |         |              |    |   |
|   |  |  Node.js 18-slim       |      |  Node.js 18-slim       |    |   |
|   |  +------------------------+      +------------------------+    |   |
|   |           |                               |                    |   |
|   +-----------|-------------------------------|--------------------+   |
|               |                               |                        |
+---------------|-------------------------------|------------------------+
                |                               |
                v                               v
        +---------------+              +-----------------+
        |   User        |              |  MongoDB Atlas  |
        |   Browser     |              |  Cloud Database |
        |   :3000       |              |  (Sound_lk)     |
        +---------------+              +-----------------+

+------------------------------------------------------------------------+
|   VOLUME MOUNTS                                                        |
|   +------------------------+                                           |
|   |  backend-uploads       | --> /app/uploads (Product Images)         |
|   +------------------------+                                           |
+------------------------------------------------------------------------+
'''
diagram3.add_run(diagram3_text)
for run in diagram3.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

doc.add_paragraph('Figure 2.3: Docker Container Architecture')

doc.add_page_break()

doc.add_heading('2.4 Component Connectivity Diagram', 2)

doc.add_paragraph(
    'This diagram details the connectivity between all application components '
    'including frontend, backend, database, and external services.'
)

diagram4 = doc.add_paragraph()
diagram4.alignment = WD_ALIGN_PARAGRAPH.CENTER
diagram4_text = '''
+============================================================================+
|                    COMPONENT CONNECTIVITY DIAGRAM                          |
+============================================================================+

   +------------------+                              +------------------+
   |     CLIENT       |                              |   ADMIN PANEL    |
   |     BROWSER      |                              |    (React)       |
   +--------+---------+                              +--------+---------+
            |                                                 |
            |  HTTP (Port 3000)                              |
            +---------------------+     +---------------------+
                                  |     |
                                  v     v
                         +------------------+
                         |    FRONTEND      |
                         |    CONTAINER     |
                         +------------------+
                         |  React 18.3.1    |
                         |  Vite 6.0.5      |
                         |  react-router    |
                         |  axios           |
                         +--------+---------+
                                  |
                                  | REST API Calls
                                  | (axios -> http://backend:5000)
                                  |
                                  v
                         +------------------+
                         |    BACKEND       |
                         |    CONTAINER     |
                         +------------------+
                         |  Express 4.21.2  |
                         |  JWT Auth        |
                         |  Multer          |
                         |  Mongoose 8.0.0  |
                         +--------+---------+
                                  |
                +-----------------|------------------+
                |                 |                  |
                v                 v                  v
      +-------------+    +---------------+    +-------------+
      | /api/auth   |    | /api/products |    | /api/orders |
      | Routes      |    | Routes        |    | Routes      |
      +-------------+    +---------------+    +-------------+
                |                 |                  |
                +-----------------|------------------+
                                  |
                                  | MongoDB Driver
                                  | (mongoose)
                                  v
                         +------------------+
                         |  MONGODB ATLAS   |
                         |  Cloud Database  |
                         +------------------+
                         |  Database:       |
                         |  Sound_lk        |
                         +------------------+
                         |  Collections:    |
                         |  - users         |
                         |  - products      |
                         |  - carts         |
                         |  - orders        |
                         +------------------+
'''
diagram4.add_run(diagram4_text)
for run in diagram4.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

doc.add_paragraph('Figure 2.4: Application Component Connectivity')

doc.add_page_break()

doc.add_heading('2.5 Diagram Explanation', 2)

doc.add_paragraph(
    'The CI/CD architecture for SoundPlus++ consists of the following key components and their interactions:'
)

doc.add_heading('Git Tools - GitHub', 3)
doc.add_paragraph(
    'GitHub serves as the central version control system for the SoundPlus++ project. '
    'The repository (https://github.com/Thiwankabanadara5400/Soundplus.git) hosts all source code, '
    'Docker configurations, and CI/CD pipeline definitions. Developers push code changes to the '
    'main branch, which triggers the automated CI/CD pipelines.'
)

doc.add_heading('CI Tool - Jenkins', 3)
doc.add_paragraph(
    'Jenkins is configured as the local CI/CD orchestrator. When code is pushed to GitHub, '
    'a webhook triggers the Jenkins pipeline defined in the Jenkinsfile. Jenkins performs '
    'code checkout, environment setup, Docker image building, and service deployment. '
    'The pipeline includes health checks to verify successful deployment.'
)

doc.add_heading('Configuration Management - Environment Variables', 3)
doc.add_paragraph(
    'Environment configuration is managed through .env files for both frontend and backend services. '
    'The Jenkins pipeline automatically creates these environment files during the Setup Environment stage, '
    'ensuring consistent configuration across deployments.'
)

doc.add_heading('Containerization - Docker', 3)
doc.add_paragraph(
    'Docker provides containerization for both frontend and backend applications. Each service '
    'has its own Dockerfile that defines the build process using Node.js 18-slim as the base image. '
    'Docker Compose orchestrates the multi-container deployment, managing networking between containers '
    'and volume mounts for persistent data storage.'
)

doc.add_heading('Container Connectivity', 3)
doc.add_paragraph(
    'The frontend and backend containers communicate over a Docker bridge network (soundplus-network). '
    'The frontend container (port 3000) makes REST API calls to the backend container (port 5000) using axios. '
    'The backend container connects to MongoDB Atlas for data persistence. All inter-service communication '
    'is secured within the Docker network.'
)

doc.add_page_break()

# Section 3: Part 2 - Automation Approach
doc.add_heading('3. Part 2: Automation Approach', 1)

doc.add_heading('3.1 DevOps Tools and Versions', 2)

doc.add_paragraph(
    'The following table describes all DevOps tools used in the SoundPlus++ deployment pipeline:'
)

devops_rows = [
    ('Git', '2.x', 'Version control system for source code management and collaboration'),
    ('GitHub', 'Cloud', 'Remote repository hosting, pull requests, and code review'),
    ('GitHub Actions', 'Latest', 'Cloud-based CI/CD pipeline for automated builds and deployments'),
    ('Jenkins', 'LTS (2.x)', 'Local CI/CD server for build automation and deployment orchestration'),
    ('Docker', '24.x', 'Containerization platform for packaging applications with dependencies'),
    ('Docker Compose', 'v2', 'Multi-container orchestration tool for defining and running services'),
    ('Docker Hub', 'Cloud', 'Container registry for storing and distributing Docker images'),
    ('Node.js', '18.x LTS', 'JavaScript runtime for running frontend and backend applications'),
    ('npm', '9.x', 'Package manager for installing and managing JavaScript dependencies'),
    ('AWS EC2', 'Cloud', 'Cloud virtual machine instances for production deployment (optional)'),
]

create_table_with_header(doc, ['Tool', 'Version', 'Purpose'], devops_rows)

doc.add_page_break()

doc.add_heading('3.2 Application Tools and Dependencies', 2)

doc.add_heading('Frontend Dependencies', 3)

frontend_rows = [
    ('react', '18.3.1', 'Core UI library for building component-based interfaces'),
    ('react-dom', '18.3.1', 'React DOM rendering engine'),
    ('react-router-dom', '7.1.1', 'Client-side routing for single-page application navigation'),
    ('axios', '1.11.0', 'HTTP client for making REST API requests to backend'),
    ('react-icons', '5.5.0', 'Icon library providing popular icon sets'),
    ('react-toastify', '11.0.5', 'Toast notification library for user feedback'),
    ('swiper', '11.1.0', 'Touch slider/carousel component for product displays'),
    ('vite', '6.0.5', 'Fast build tool and development server'),
    ('@vitejs/plugin-react', '4.3.4', 'Vite plugin for React support with Fast Refresh'),
    ('eslint', '9.17.0', 'JavaScript linter for code quality enforcement'),
]

create_table_with_header(doc, ['Package', 'Version', 'Purpose'], frontend_rows)

doc.add_paragraph()
doc.add_heading('Backend Dependencies', 3)

backend_rows = [
    ('express', '4.21.2', 'Web framework for building REST API server'),
    ('mongoose', '8.0.0', 'MongoDB ODM for database modeling and queries'),
    ('bcrypt', '5.1.1', 'Password hashing library for secure authentication'),
    ('jsonwebtoken', '9.0.2', 'JWT implementation for token-based authentication'),
    ('cors', '2.8.5', 'Middleware for enabling Cross-Origin Resource Sharing'),
    ('body-parser', '1.20.3', 'Middleware for parsing request bodies'),
    ('cookie-parser', '1.4.7', 'Middleware for parsing cookies'),
    ('multer', '1.4.5-lts.1', 'Middleware for handling file uploads'),
    ('dotenv', '16.4.7', 'Environment variable management'),
    ('nodemon', '3.1.9', 'Development tool for auto-reloading on file changes'),
]

create_table_with_header(doc, ['Package', 'Version', 'Purpose'], backend_rows)

doc.add_page_break()

doc.add_heading('Database Configuration', 3)

db_rows = [
    ('MongoDB Atlas', 'Cloud', 'Cloud-hosted MongoDB database service'),
    ('Database Name', 'Sound_lk', 'Main application database'),
    ('Connection', 'MongoDB Driver', 'mongoose ODM for Node.js'),
]

create_table_with_header(doc, ['Component', 'Value', 'Description'], db_rows)

doc.add_paragraph()
doc.add_heading('Database Collections', 3)

collections_rows = [
    ('users', 'username, email, password, role, createdAt', 'User authentication and profiles'),
    ('products', 'name, price, category, brand, features, etc.', 'Product catalog information'),
    ('carts', 'userId, productId, quantity', 'Shopping cart items'),
    ('orders', 'userId, items, totalAmount, status, shippingAddress', 'Order records'),
]

create_table_with_header(doc, ['Collection', 'Fields', 'Purpose'], collections_rows)

doc.add_page_break()

doc.add_heading('3.3 Jenkins Pipeline Stages', 2)

doc.add_paragraph(
    'The Jenkins pipeline (Jenkinsfile) automates the deployment process through the following stages:'
)

pipeline_rows = [
    ('1', 'Checkout', 'Clones the source code from GitHub repository (main branch)'),
    ('2', 'Pre-flight Check', 'Validates Docker and Docker Compose versions, cleans up existing containers'),
    ('3', 'Setup Environment', 'Creates .env files with required configuration variables'),
    ('4', 'Build Images', 'Builds Docker images for frontend and backend using docker-compose build'),
    ('5', 'Start Services', 'Launches containers using docker-compose up in detached mode'),
    ('6', 'Verify Services', 'Performs health checks on backend /health endpoint'),
    ('7', 'Success', 'Displays deployment information and access URLs'),
]

create_table_with_header(doc, ['Stage', 'Name', 'Description'], pipeline_rows)

doc.add_paragraph()
doc.add_heading('Jenkins Pipeline Configuration', 3)

config_rows = [
    ('COMPOSE_PROJECT_NAME', 'soundplus', 'Docker Compose project identifier'),
    ('PROJECT_NAME', 'SoundPlus++', 'Display name for the project'),
    ('SCM Repository', 'GitHub', 'https://github.com/Thiwankabanadara5400/Soundplus.git'),
    ('Branch', 'main', 'Default branch for deployment'),
]

create_table_with_header(doc, ['Parameter', 'Value', 'Description'], config_rows)

doc.add_page_break()

doc.add_heading('3.4 GitHub Actions Pipeline', 2)

doc.add_paragraph(
    'GitHub Actions provides cloud-based CI/CD with the following workflow configuration:'
)

ga_rows = [
    ('backend-build', 'Build', 'Builds backend with Node.js 18, installs dependencies'),
    ('frontend-build', 'Build', 'Builds frontend with Node.js 18, installs dependencies'),
    ('docker-push', 'Push', 'Builds and pushes Docker images to Docker Hub registry'),
    ('deploy', 'Deploy', 'Deploys to AWS EC2 instance via SSH'),
]

create_table_with_header(doc, ['Job Name', 'Type', 'Description'], ga_rows)

doc.add_paragraph()
doc.add_heading('GitHub Actions Triggers', 3)
doc.add_paragraph(
    '- Push events to main/master branches\n'
    '- Pull request events to main/master branches\n'
    '- Docker push and deploy jobs only run on main/master branch pushes'
)

doc.add_heading('Required GitHub Secrets', 3)

secrets_rows = [
    ('DOCKER_USERNAME', 'Docker Hub authentication username'),
    ('DOCKER_PASSWORD', 'Docker Hub authentication password/token'),
    ('AWS_ACCESS_KEY_ID', 'AWS IAM access key for EC2 deployment'),
    ('AWS_SECRET_ACCESS_KEY', 'AWS IAM secret key for EC2 deployment'),
    ('AWS_REGION', 'AWS region for EC2 instance'),
    ('EC2_SSH_PRIVATE_KEY', 'SSH private key for EC2 access'),
    ('EC2_HOST', 'EC2 instance hostname or IP address'),
    ('EC2_USER', 'SSH username for EC2 instance'),
]

create_table_with_header(doc, ['Secret Name', 'Purpose'], secrets_rows)

doc.add_page_break()

doc.add_heading('3.5 Deployment Automation Flow', 2)

doc.add_paragraph(
    'The complete deployment automation follows this sequence:'
)

flow_steps = [
    ('1', 'Developer pushes code changes to GitHub repository'),
    ('2', 'GitHub webhook triggers Jenkins pipeline OR GitHub Actions workflow'),
    ('3', 'Pipeline clones repository and validates environment'),
    ('4', 'Environment variables are configured from templates'),
    ('5', 'Docker images are built for frontend and backend'),
    ('6', 'Images are tagged and pushed to Docker Hub registry'),
    ('7', 'Docker Compose starts containers on target environment'),
    ('8', 'Health checks verify service availability'),
    ('9', 'Deployment status is reported (success/failure)'),
    ('10', 'Application is accessible at configured ports'),
]

create_table_with_header(doc, ['Step', 'Action'], flow_steps)

doc.add_paragraph()
doc.add_heading('Automation Scripts', 3)

scripts_rows = [
    ('docker-rebuild.sh', 'Complete Docker rebuild with cache cleanup'),
    ('docker-push.sh', 'Push images to Docker Hub with proper tagging'),
    ('docker-check.sh', 'Health check validation for running services'),
    ('scripts/setup-jenkins.sh', 'Automated Jenkins server configuration'),
]

create_table_with_header(doc, ['Script', 'Purpose'], scripts_rows)

doc.add_page_break()

# Section 4: Environment Configuration
doc.add_heading('4. Environment Configuration', 1)

doc.add_heading('Backend Environment Variables', 2)

backend_env_rows = [
    ('PORT', '5000', 'Backend server port'),
    ('NODE_ENV', 'development/production', 'Runtime environment mode'),
    ('MONGODB_URI', 'mongodb+srv://...', 'MongoDB Atlas connection string'),
    ('DB_NAME', 'Sound_lk', 'Database name'),
    ('JWT_SECRET', 'soundplus_secret_key_2025', 'Secret for JWT token signing'),
    ('CORS_ORIGIN', 'http://localhost:3000', 'Allowed CORS origin'),
]

create_table_with_header(doc, ['Variable', 'Value', 'Description'], backend_env_rows)

doc.add_paragraph()
doc.add_heading('Frontend Environment Variables', 2)

frontend_env_rows = [
    ('VITE_API_URL', 'http://localhost:5000', 'Backend API base URL'),
]

create_table_with_header(doc, ['Variable', 'Value', 'Description'], frontend_env_rows)

doc.add_paragraph()
doc.add_heading('Docker Compose Configuration', 2)

docker_rows = [
    ('Frontend Container', 'soundplus-frontend', 'Port 3000'),
    ('Backend Container', 'soundplus-backend', 'Port 5000'),
    ('Network', 'soundplus-network', 'Bridge driver'),
    ('Volume', 'backend-uploads', 'Product image storage'),
]

create_table_with_header(doc, ['Component', 'Name', 'Configuration'], docker_rows)

doc.add_page_break()

# Section 5: Security Considerations
doc.add_heading('5. Security Considerations', 1)

doc.add_paragraph(
    'The SoundPlus++ application implements several security measures:'
)

security_rows = [
    ('Authentication', 'JWT tokens with 7-day expiration, stored in httpOnly cookies'),
    ('Password Security', 'bcrypt hashing with 10 salt rounds'),
    ('Access Control', 'Role-based access (user/admin) for protected routes'),
    ('CORS', 'Restricted to configured origins only'),
    ('File Uploads', '10MB limit, image format validation (jpeg, png, gif, webp)'),
    ('Environment', 'Sensitive data stored in .env files, not in code'),
    ('Network', 'Docker bridge network isolates container communication'),
]

create_table_with_header(doc, ['Security Feature', 'Implementation'], security_rows)

doc.add_page_break()

# Section 6: Conclusion
doc.add_heading('6. Conclusion', 1)

doc.add_paragraph(
    'The SoundPlus++ project implements a comprehensive CI/CD pipeline that ensures '
    'reliable and efficient software delivery. The architecture combines local Jenkins '
    'pipelines with cloud-based GitHub Actions to provide flexibility in deployment options.'
)

doc.add_paragraph(
    'Key highlights of the automation approach include:'
)

highlights = doc.add_paragraph()
highlights.add_run('\n- Fully containerized application using Docker and Docker Compose')
highlights.add_run('\n- Dual CI/CD options: Jenkins (local) and GitHub Actions (cloud)')
highlights.add_run('\n- Automated environment configuration and health checks')
highlights.add_run('\n- Secure container networking with isolated communication')
highlights.add_run('\n- Scalable architecture supporting multiple deployment targets')
highlights.add_run('\n- Comprehensive monitoring through health check endpoints')

doc.add_paragraph()
doc.add_paragraph(
    'This documentation provides a complete overview of the CI/CD design and automation '
    'approach for the SoundPlus++ e-commerce platform, enabling consistent and repeatable '
    'deployments across development, staging, and production environments.'
)

# Appendix
doc.add_page_break()
doc.add_heading('Appendix A: Quick Reference', 1)

doc.add_heading('Access URLs', 2)
urls_rows = [
    ('Frontend', 'http://localhost:3000', 'User interface'),
    ('Backend API', 'http://localhost:5000', 'REST API endpoints'),
    ('Health Check', 'http://localhost:5000/health', 'Backend health status'),
]
create_table_with_header(doc, ['Service', 'URL', 'Description'], urls_rows)

doc.add_paragraph()
doc.add_heading('Docker Commands', 2)
commands_rows = [
    ('docker-compose up --build', 'Build and start all services'),
    ('docker-compose down', 'Stop and remove all containers'),
    ('docker-compose logs -f', 'View real-time logs'),
    ('docker-compose ps', 'List running containers'),
]
create_table_with_header(doc, ['Command', 'Description'], commands_rows)

doc.add_paragraph()
doc.add_heading('Repository Information', 2)
repo_rows = [
    ('GitHub URL', 'https://github.com/Thiwankabanadara5400/Soundplus.git'),
    ('Default Branch', 'main'),
    ('License', 'ISC'),
]
create_table_with_header(doc, ['Item', 'Value'], repo_rows)

# Save document
doc.save('D:/Docker project/SoundPlus++/SoundPlus_CICD_Documentation.docx')
print('Document created successfully: SoundPlus_CICD_Documentation.docx')
